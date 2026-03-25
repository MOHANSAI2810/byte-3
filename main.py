import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import time
import json
import random
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.data_loader import load_and_process_data
from datetime import datetime
import re

# Import for spider chart
import numpy as np
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
from io import BytesIO


# --- CONFIGURATION ---
DEEPSEEK_API_KEY = "sk-5fc2fb744dac49fd821977bdfea028b2"
INPUT_FOLDER = "input"
OUTPUT_FOLDER = "output"
EXPLAINABILITY_FOLDER = "explainability_logs"

# --- MODEL: DEEPSEEK CHAT ---
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

# =========================
# INTEREST LEVEL LOGIC
# =========================
INTEREST_LEVEL_SYNONYMS = {
    "HIGH": [
        "Highly Developed Interest",
        "Strong Mastery and Enthusiasm",
        "Advanced Engagement and Confidence",
        "Consistent Strength and Passion"
    ],
    "STRONG": [
        "Showing Strong Interest",
        "Positive and Active Engagement",
        "Well-Developed Interest",
        "Growing Strength in This Area"
    ],
    "DEVELOPING": [
        "Developing Engagement",
        "Emerging Interest",
        "Building Confidence Gradually",
        "Progressing With Guidance"
    ],
    "EXPLORING": [
        "Exploring Preferences",
        "Early Curiosity Observed",
        "Initial Engagement Noted",
        "Beginning to Show Interest"
    ],
    "EARLY": [
        "Early Exposure",
        "Limited Experience So Far",
        "Initial Familiarity Developing",
        "Requires More Opportunities"
    ],
    "NA": ["Not Assessed Yet"]
}

def get_interest_label(score):
    if score is None:
        return random.choice(INTEREST_LEVEL_SYNONYMS["NA"])
    
    if score >= 85:
        return random.choice(INTEREST_LEVEL_SYNONYMS["HIGH"])
    elif score >= 60:
        return random.choice(INTEREST_LEVEL_SYNONYMS["STRONG"])
    elif score >= 40:
        return random.choice(INTEREST_LEVEL_SYNONYMS["DEVELOPING"])
    elif score >= 20:
        return random.choice(INTEREST_LEVEL_SYNONYMS["EXPLORING"])
    else:
        return random.choice(INTEREST_LEVEL_SYNONYMS["EARLY"])

# =========================
# DOCX TABLE STYLING HELPERS
# =========================
def style_table_header(row, bg_color="1F5E78", text_color=RGBColor(255, 255, 255)):
    """Applies background color and white bold text to a header row."""
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg_color)
        tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = text_color

def set_column_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

def add_formatted_text_with_bold(cell, text):
    """
    Adds text to a cell with **bold** markdown-style formatting.
    Converts **text** into actual bold runs.
    """
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()  # Clear existing content
    
    # Split by ** markers
    parts = text.split('**')
    
    for i, part in enumerate(parts):
        if not part:  # Skip empty parts
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(9)
        # Odd indices are between ** markers (should be bold)
        if i % 2 == 1:
            run.bold = True

def format_single_point(cell, text):
    """
    Formats text to show only the first point from a numbered list.
    Extracts just the first item and removes any extra numbering.
    """
    # Clear existing content
    cell._element.clear_content()
    
    # Extract just the first point (before the first semicolon)
    if ';' in text:
        first_point = text.split(';')[0].strip()
    else:
        first_point = text.strip()
    
    # Remove any leading numbers like "1." or "1" if they exist
    # This pattern matches: optional number, optional period, optional space at the beginning
    cleaned_point = re.sub(r'^\d+\.?\s*', '', first_point)
    
    # Add as a single paragraph
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    
    # Add the cleaned text (without the number)
    run = p.add_run(cleaned_point)
    run.font.size = Pt(9)

def create_bordered_paragraph(cell, text, bold=False):
    """
    Creates a bordered paragraph cell for the Notes section.
    """
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    
    # Add border to the cell
    tc_pr = cell._tc.get_or_add_tcPr()
    
    # Add borders to the cell
    for border_side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_pr.append(border)

# =========================
# SPIDER CHART GENERATION - UPDATED
# =========================
def create_spider_chart(student_data, class_avg_data):
    """
    Creates a spider/radar chart comparing student performance vs class average.
    Returns a BytesIO object containing the PNG image.
    Updated with:
    - Labels moved away from circle
    - Transparent background
    - Cleaner appearance
    """
    try:
        # Extract activity names and scores
        activities = []
        student_scores = []
        class_scores = []
        
        for activity, details in student_data["activities"].items():
            # Skip "Others" with 0% or None
            score = details.get("average")
            if activity.lower() == "others" and (score is None or score == 0):
                continue
                
            activities.append(activity)
            student_scores.append(round(score) if score is not None else 0)
            
            # Get class average for this activity - handle both dict and direct value
            class_item = class_avg_data.get(activity, 0)
            if isinstance(class_item, dict):
                class_score = class_item.get("average", 0)
            else:
                class_score = class_item  # It's already a number
            class_scores.append(round(class_score) if class_score else 0)
        
        # If no activities or too few, return None
        if len(activities) < 2:
            print("   ⚠️ Not enough activities for spider chart (need at least 2)")
            return None
        
        num_vars = len(activities)
        
        # Compute angle for each axis
        angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
        
        # Close the plot (repeat first value)
        student_scores_plot = student_scores + student_scores[:1]
        class_scores_plot = class_scores + class_scores[:1]
        angles_plot = angles + angles[:1]
        
        # Create figure with transparent background
        fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
        fig.patch.set_alpha(0)  # Transparent figure background
        ax.set_facecolor('none')  # Transparent axes background
        
        # Set rotation
        ax.set_theta_offset(np.pi / 2)
        ax.set_theta_direction(-1)
        
        # Draw axis labels - MOVED AWAY FROM CIRCLE (increased padding)
        ax.set_xticks(angles)
        ax.set_xticklabels(activities, size=9, weight='bold')
        
        # Increase label padding to move them away from the circle
        ax.tick_params(axis='x', pad=25)  # Increased padding to 25 points (approximately 0.35 inches / 0.9 cm)
        
        # Set radial scale
        ax.set_rlabel_position(0)
        ax.set_yticks([25, 50, 75, 100])
        ax.set_yticklabels(['25', '50', '75', '100'], size=8, color='gray')
        ax.set_ylim(0, 100)
        
        # Make grid lines more prominent
        ax.grid(True, linestyle='--', color='gray', alpha=0.5, linewidth=1.2)
        
        # Thicken the outer circular border (the spine)
        ax.spines['polar'].set_linewidth(2.0)
        ax.spines['polar'].set_color('#333333')
        
        # Plot student average (blue)
        ax.plot(angles_plot, student_scores_plot, linewidth=2.5, marker='o', 
                label='Student average', color='#4472C4', markersize=6)
        ax.fill(angles_plot, student_scores_plot, alpha=0.25, color='#4472C4')
        
        # Plot class average (red/orange)
        ax.plot(angles_plot, class_scores_plot, linewidth=2.5, marker='o', 
                label='Class average', color='#ED7D31', markersize=6)
        ax.fill(angles_plot, class_scores_plot, alpha=0.15, color='#ED7D31')
        
        # Legend with transparent background
        legend = ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.1), 
                          frameon=False, fontsize=9)
        
        # Save to BytesIO with transparent background
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=200, bbox_inches='tight', 
                   transparent=True, pad_inches=0.1)
        img_buffer.seek(0)
        
        return img_buffer
        
    except Exception as e:
        print(f"   ⚠️ Error creating spider chart: {e}")
        return None
    finally:
        # Always close the figure to free memory
        plt.close('all')

# =========================
# LLM ANALYSIS
# =========================
def generate_llm_analysis(student_data, student_name, class_name):
    # 1. Prepare Data with Calculated Synonyms
    raw_student_data = "Activity | Score | Interest Level | Feedback\n"
    for activity, details in student_data["activities"].items():
        score = round(details["average"]) if details.get("average") is not None else None
        
        interest_label = get_interest_label(score)
        
        score_text = f"{score}%" if score is not None else "-"
        feedbacks = "; ".join(details["feedbacks"][:5]) if details.get("feedbacks") else "No specific feedback recorded"
        raw_student_data += f"{activity} | {score_text} | {interest_label} | {feedbacks}\n"

    # 2. Tone Block Logic
    tone_block = "BLOCK P – Primary Classes 3–4"
    try:
        class_num = int(re.search(r"\d+", str(class_name)).group())
        if class_num in [3, 4]:
            tone_block = "BLOCK P – PRIMARY (Classes 3–4)"
        elif class_num in [5, 6]:
            tone_block = "BLOCK U – UPPER PRIMARY (Classes 5–6)"
        elif class_num in [7, 8, 9]:
            tone_block = "BLOCK M – MIDDLE SCHOOL (Classes 7–9)"
    except:
        pass

    # 3. Prompt (UPDATED - subject_career_combinations still needed for data but won't be displayed)
    prompt = f"""
<SYSTEM PROMPT – PARENT CO-CURRICULAR REPORT>

You are an expert Educational Consultant designing parent-friendly,
practical, and class-appropriate co-curricular career recommendations
for Indian families.

Your goal is to help parents clearly understand:
• What the child is naturally inclined towards
• What practical steps parents can take next
• What future directions may gradually open up (without pressure)

LANGUAGE & STYLE RULES (STRICT)
• Use only simple, everyday words
• Write warmly and encouragingly
• Never use the student's name
• Never use pronouns (he / she / they)
• Always write in present tense
• Write for Indian parents (realistic, grounded)

CONTEXT
Class Level: {class_name}
Tone & Career Block: {tone_block}

RAW STUDENT DATA (pipe-delimited):
{raw_student_data}

-------------------------------------------------
CRITICAL RULES
-------------------------------------------------

1) INTEREST LEVELS (NON-NEGOTIABLE)
• Interest levels are already provided
• Use the exact wording from the data
• Do NOT rephrase, explain, or substitute

-------------------------------------------------

2) OBSERVATIONS (MANDATORY)
• Every subject and activity in the data must appear
• Write 1–2 simple sentences per area
• Focus on effort, engagement, curiosity, and learning behavior
• Do NOT compare with other students

-------------------------------------------------

-------------------------------------------------
3) RECOMMENDATIONS (TOP 2 ONLY – WITH PARENT ACTIONABILITY)

Select ONLY the top 2 highest engagement areas.

-------------------------------------------------
3.1 ACTIVITY-ONLY RULE
-------------------------------------------------
• Recommendations MUST be based ONLY on the specific activity name (e.g., "Computers", "Sports", "Arts").
• Do NOT combine activities with Academics in the first column or title of the recommendation.
• Even if Academics appears in the top scores, treat it as its own area or prioritize the next highest non-academic activity if "Academic + Activity" was previously requested.
• Clearly explain how the activity supports overall learning in the description, but keep the header focused on the activity itself.

-------------------------------------------------
3.2 CLASS-BASED CAREER GUIDANCE
-------------------------------------------------

For Classes 3–5:
• Do NOT mention job titles
• Focus on exposure, confidence, habits, and skill discovery
• Career direction must remain broad and future-facing

For Classes 6–7:
• Introduce career fields (not final roles)
• Explain how current interests connect to these fields
• Suggest structured learning, clubs, or guided practice

For Classes 8–9:
• Mention specific career paths where relevant
• Keep tone supportive and non-pressurizing
• Emphasize readiness, not final decisions

-------------------------------------------------
3.3 FORMATTING REQUIREMENTS (STRICT)
-------------------------------------------------

CAREER PATHWAYS FORMAT:
• Start with a brief intro sentence (1 line)
• List 3-5 specific career paths in BOLD using **Career Name** format
• Keep descriptions practical and realistic
• Connect careers logically to the student's interests

Example Format:
"This well-developed interest can lead to diverse career paths in sports. Exploring **coaching**, **sports management**, **sports journalism**, or even **professional athleticism** is possible. Developing strong teamwork and leadership skills also benefits many other fields."

-------------------------------------------------

PARENT ACTIONABILITY RULE (MANDATORY):
-------------------------------------------------

For EVERY recommendation, provide EXACTLY 3 parent actions in numbered format.

Each action MUST:
• Be specific and practical (not generic motivation)
• Appropriate for the class level
• Directly connected to the student's interest and engagement data
• Easy for Indian parents to act on within the next 6–12 months
• Mention the type of class, activity, or routine
• Clearly explain what skill the action strengthens
• Avoid vague phrases like "encourage interest", "explore more", "provide exposure"

NUMBERED FORMAT (MANDATORY):
Write actions as: "1. [Action]; 2. [Action]; 3. [Action]"

Example Format:
"1. Enroll in a structured badminton academy or shuttle coaching program to strengthen specific skills; 2. Encourage participation in school or inter-school sports competitions to build competitive spirit and teamwork; 3. Provide opportunities to learn about sports nutrition and fitness routines to support physical development."

CLASS-WISE EXPECTATIONS:

For Classes 3–5:
• Focus on habit building, confidence, and exposure
• Suggest hobby classes, simple home routines, and school-level participation
• Avoid exams, certifications, or career pressure

For Classes 6–7:
• Focus on structured learning and skill strengthening
• Suggest guided classes, small projects, competitions, or clubs
• Introduce career fields indirectly

For Classes 8–9:
• Focus on preparation and clarity
• Suggest foundation courses, portfolios, mentoring, or academies
• Clearly connect actions to future career paths

-------------------------------------------------
5) HOW PARENTS CAN HELP: PARENT TIPS (REPLACEMENT)
-------------------------------------------------
Provide exactly 3 high-quality tips in the "parent_tips" field. 
Use this specific style (combining activity context with practical timing):

Style Examples:
1. Weekend Exploration (Computers): Dedicate a weekend afternoon to a fun, structured online lesson that teaches basic data organization or typing skills, connecting it to the interest in data entry.
2. Short Daily Practice (Drawing): Encourage 15 minutes of sketching simple objects or geometric shapes in the evening to maintain the interest in exact measurements and shading techniques.
3. Encourage and Talk: Dedicate 10–15 minutes weekly to talk about what excites your child and offer your appreciation for their efforts in these activities.

Requirements for the 3 tips:
- Tip 1: Should be a "Weekend Exploration" related to the child's top interest.
- Tip 2: Should be a "Short Daily Practice" (15 mins) related to the child's second interest.
- Tip 3: Should be a "Encourage and Talk" tip (10-15 mins weekly).

-------------------------------------------------
6) CONCLUSION
• Write a warm, reassuring summary
• Emphasize steady growth and natural strengths
• Reinforce that exploration is healthy at this stage
• Do NOT mention the class name directly

-------------------------------------------------
CAREER COMBINATION KNOWLEDGE BASE
-------------------------------------------------

Use the following subject + activity combinations to populate the "subject_career_combinations" field.
Only include combinations when BOTH the activity AND a related academic subject show strong engagement.

SPORTS COMBINATIONS:
• Sports + Science: Sports Scientist, Physiotherapist, Sports Nutritionist, Strength & Conditioning Coach, Exercise Physiologist
• Sports + Maths: Sports Data Analyst, Performance Analyst, Sports Statistics Expert
• Sports + Languages: Sports Journalist, Sports Commentator, Sports Content Creator, Sports PR & Media Manager
• Sports + Social: Sports Administrator, Sports Policy Analyst, Sports Management Professional, Community Sports Officer

LIBRARY (READING/BOOKS) COMBINATIONS:
• Library + Languages: Author/Writer, Editor/Publisher, Journalist, Scriptwriter, Content Strategist
• Library + Social: Historian, Civil Services Officer, Political Analyst, Policy Researcher, Think-Tank Researcher
• Library + Science: Research Scholar, Academic Scientist, Science Communicator, Science Writer
• Library + Maths: Economist, Actuary, Data Research Analyst, Financial Analyst

COMPUTERS COMBINATIONS:
• Computers + Maths: Software Engineer, Data Scientist, AI/ML Engineer, Game Developer, Quant Analyst
• Computers + Science: Robotics Engineer, Bioinformatics Specialist, Environmental Tech Analyst, Health-Tech Developer
• Computers + Languages: Technical Writer, UX Writer, Product Manager, Digital Marketer, Instructional Designer
• Computers + Social: Civic-Tech Specialist, GovTech Consultant, Digital Policy Analyst, Cyber Law Professional

ARTS (Drawing, Music, Dance, Design) COMBINATIONS:
• Arts + Languages: Filmmaker, Lyricist, Screenwriter, Theatre Artist, Creative Director
• Arts + Social: Fashion Designer, Cultural Researcher, Interior Designer, Heritage Conservationist
• Arts + Science: Medical Illustrator, Industrial Designer, Product Designer
• Arts + Maths: Architect, UI/UX Designer, Animation & VFX Artist, Game Designer

MUSIC/DANCE COMBINATIONS:
• Music + Languages: Lyricist, Music Journalist, Radio Jockey, Podcast Host, Music Educator
• Music + Maths: Sound Engineer, Music Producer, Audio Technology Specialist
• Dance + Science: Dance Therapist, Movement Analyst, Fitness Choreographer
• Dance + Social: Cultural Program Coordinator, Dance Historian, Arts Administrator

LOGIC FOR POPULATING "subject_career_combinations":
1. Check if the student has high engagement in BOTH an activity (Sports, Library, Computers, Arts, Music, Dance) AND a related academic subject (Science, Maths, Languages, Social)
2. If YES → Select 4-5 best-fit careers from the matching combination above
3. If NO clear academic pairing → Leave this field as null or empty string
4. Format as plain text list separated by commas (no bold, no bullets)
5. Example output: "Sports Scientist, Physiotherapist, Sports Nutritionist, Strength & Conditioning Coach"

-------------------------------------------------
OUTPUT FORMAT (JSON ONLY)
-------------------------------------------------
{{
  "observations": [
    {{
      "activity": "Name",
      "engagement": 90,
      "interest_level": "Exact String From Data",
      "observation": "Text"
    }}
  ],
  "recommendations": [
    {{
      "area": "Name",
      "engagement": 90,
      "future_pathways": "Brief intro sentence. Mention **Career1**, **Career2**, **Career3** in bold. Additional context about skills.",
      "subject_career_combinations": "Career1, Career2, Career3, Career4, Career5 (only if academic subject pairing exists, otherwise null)",
      "support_activities": "1. First specific action with details; 2. Second specific action with details; 3. Third specific action with details"
    }}
  ],
  "parent_tips": ["Tip 1", "Tip 2", "Tip 3"],
  "conclusion": "Text"
}}
"""
    
    # 4. DeepSeek Completion (With JSON Enforcement)
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "You are an expert Educational Consultant. Always respond in valid JSON format."},
                    {"role": "user", "content": prompt}
                ],
                response_format={'type': 'json_object'}
            )
            text = response.choices[0].message.content.strip()
            
            parsed = json.loads(text)
            return parsed, {"input": raw_student_data}

        except Exception as e:
            print(f"   > API Error (Attempt {attempt+1}): {e}")
            if "429" in str(e):
                print("   > Quota exceeded. Waiting 60s to reset...")
                time.sleep(60) 
            else:
                time.sleep(5)
    return None, None

# =========================
# WORD REPORT CREATION
# =========================
def create_word_doc(student_name, analysis, class_name, student_data, class_avg_data, output_dir=None):
    doc = Document()
    
    # Use specified output_dir or fallback to global constant
    target_folder = output_dir if output_dir else OUTPUT_FOLDER
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # ========================================
    # PAGE 1: STUDENT INFO, PURPOSE, KEY INSIGHTS, CHART, OBSERVATIONS
    # ========================================
    
    # Student Info Table (compact)
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    set_table_borders(info_table)
    set_column_widths(info_table, [Inches(1.2), Inches(5.8)])
    
    # Row 1: Name
    info_table.rows[0].cells[0].text = "Name"
    info_table.rows[0].cells[1].text = student_name
    style_table_header(info_table.rows[0], bg_color="1F5E78")
    
    # Row 2: Class
    info_table.rows[1].cells[0].text = "Class"
    info_table.rows[1].cells[1].text = class_name
    
    # Make labels bold
    for row in info_table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Purpose Section
    purpose_heading = doc.add_heading("Purpose:", level=2)
    purpose_heading.style.font.size = Pt(11)
    purpose_heading.style.font.bold = True
    purpose_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    purpose_heading.paragraph_format.space_before = Pt(6)
    purpose_heading.paragraph_format.space_after = Pt(2)
    
    purpose_text = doc.add_paragraph(
        '"Skill Assessment Program - Building Strengths Beyond Academics "\n'
        'This report shows your child\'s interests and abilities in different school activities. '
        'With your (parents) support and our guidance, these strengths can grow into useful life skills. '
        'It helps parents notice early signs of talent and areas where extra care or practice can make a big difference.'
    )
    purpose_text.style.font.size = Pt(9)
    purpose_text.paragraph_format.space_after = Pt(4)
    
    # Key Insights Section
    insights_heading = doc.add_heading("Key Insights:", level=2)
    insights_heading.style.font.size = Pt(11)
    insights_heading.style.font.bold = True
    insights_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    insights_heading.paragraph_format.space_before = Pt(4)
    insights_heading.paragraph_format.space_after = Pt(2)
    
    insights_text = doc.add_paragraph(
        'The graph below shows how your child is doing in different areas. It tells us what they like more and where we can '
        'help them do even better. Blue represents your child, and red represents the class average.'
    )
    insights_text.style.font.size = Pt(9)
    insights_text.paragraph_format.space_after = Pt(2)
    
    # INSERT SPIDER CHART
    chart_img = create_spider_chart(student_data, class_avg_data)
    if chart_img:
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_para.paragraph_format.space_before = Pt(0)
        chart_para.paragraph_format.space_after = Pt(4)
        run = chart_para.add_run()
        run.add_picture(chart_img, width=Inches(4.5))
    
    # OBSERVATIONS SECTION
    obs_heading = doc.add_heading("Observations:", level=2)
    obs_heading.style.font.size = Pt(11)
    obs_heading.style.font.bold = True
    obs_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    obs_heading.paragraph_format.space_before = Pt(4)
    obs_heading.paragraph_format.space_after = Pt(2)
    
    subtitle = doc.add_paragraph("The table gives an overview of child's involvement and performance across co-curricular areas.")
    subtitle.style.font.size = Pt(9)
    subtitle.paragraph_format.space_after = Pt(2)
    
    # Sort observations by engagement score (Descending)
    sorted_obs = sorted(analysis.get("observations", []), 
                       key=lambda x: x.get('engagement', 0) if isinstance(x.get('engagement'), (int, float)) else -1, 
                       reverse=True)
    
    # Filter out "Others" with 0% or None
    filtered_obs = []
    for obs in sorted_obs:
        if obs.get("activity", "").strip().lower() == "others":
            eng = obs.get("engagement")
            if eng is None or eng == 0:
                continue
        filtered_obs.append(obs)
    
    # Create observations table
    obs_table = doc.add_table(rows=1, cols=4)
    obs_table.style = 'Table Grid'
    set_table_borders(obs_table)
    set_column_widths(obs_table, [Inches(1.5), Inches(0.7), Inches(1.6), Inches(3.2)])

    hdr = obs_table.rows[0]
    hdr.cells[0].text = "Interest Area"
    hdr.cells[1].text = "%"
    hdr.cells[2].text = "Interest Level"
    hdr.cells[3].text = "Key Observations"
    style_table_header(hdr, bg_color="1F5E78")

    for obs in filtered_obs:
        row = obs_table.add_row().cells
        row[0].text = obs.get("activity", "")
        eng = obs.get("engagement")
        row[1].text = f"{eng}%" if isinstance(eng, (int, float)) else "-"
        row[2].text = obs.get("interest_level", "")
        row[3].text = obs.get("observation", "")
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # PAGE BREAK - NOW PAGE 1 IS COMPLETE WITH ALL OBSERVATIONS
    doc.add_page_break()
    
    # ========================================
    # PAGE 2: RECOMMENDATIONS, TIPS, CONCLUSION, NOTES
    # ========================================
    
    # RECOMMENDATIONS SECTION
    rec_heading = doc.add_heading("Recommendations & Growth Opportunities:", level=2)
    rec_heading.style.font.size = Pt(11)
    rec_heading.style.font.bold = True
    rec_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    rec_heading.paragraph_format.space_before = Pt(4)
    rec_heading.paragraph_format.space_after = Pt(2)
    
    rec_subtitle = doc.add_paragraph("Here are some ways to turn your child's interests into long-term strengths:")
    rec_subtitle.style.font.size = Pt(9)
    rec_subtitle.paragraph_format.space_after = Pt(2)
    
    # Create Recommendations Table - NOW WITH 3 COLUMNS
    rec_table = doc.add_table(rows=1, cols=3)
    rec_table.style = 'Table Grid'
    set_table_borders(rec_table)
    set_column_widths(rec_table, [Inches(1.8), Inches(3.2), Inches(3.0)])

    # Header Row with 3 columns
    hdr = rec_table.rows[0]
    hdr.cells[0].text = "Strong Interest Area"
    hdr.cells[1].text = "Potential Careers Examples"
    hdr.cells[2].text = "Simple Activities to Support"
    style_table_header(hdr, bg_color="1F5E78")

    # Ensure we only show top 2 recommendations
    recs = analysis.get("recommendations", [])[:2]

    for rec in recs:
        row = rec_table.add_row().cells
        
        # Column 1: Area with percentage
        eng = rec.get("engagement")
        score_txt = f"\n({eng}%)" if isinstance(eng, (int, float)) else ""
        row[0].text = f"{rec.get('area', '')}{score_txt}"
        
        # Column 2: Future Pathways (with bold formatting)
        future_pathways = rec.get("future_pathways", "")
        add_formatted_text_with_bold(row[1], future_pathways)
        
        # Column 3: Support Activities - Fixed to show only first point without extra text
        support_activities = rec.get("support_activities", "")
        format_single_point(row[2], support_activities)
        
        # Set font size for columns 1
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    
    # Small space
    space_para = doc.add_paragraph()
    space_para.paragraph_format.space_after = Pt(4)
    
    # HOW PARENTS CAN HELP SECTION
    tips_heading = doc.add_heading("How Parents Can Help:", level=2)
    tips_heading.style.font.size = Pt(11)
    tips_heading.style.font.bold = True
    tips_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    tips_heading.paragraph_format.space_before = Pt(4)
    tips_heading.paragraph_format.space_after = Pt(2)
    
    for i, tip in enumerate(analysis.get("parent_tips", []), 1):
        tip_para = doc.add_paragraph(f"{i}. {tip}")
        tip_para.style.font.size = Pt(9)
        tip_para.paragraph_format.left_indent = Inches(0.2)
        tip_para.paragraph_format.space_after = Pt(3)
    
    # Small space
    space_para = doc.add_paragraph()
    space_para.paragraph_format.space_after = Pt(4)
    
    # CONCLUSION SECTION
    conclusion_heading = doc.add_heading("Conclusion:", level=2)
    conclusion_heading.style.font.size = Pt(11)
    conclusion_heading.style.font.bold = True
    conclusion_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    conclusion_heading.paragraph_format.space_before = Pt(4)
    conclusion_heading.paragraph_format.space_after = Pt(2)
    
    conclusion_para = doc.add_paragraph(analysis.get("conclusion", ""))
    conclusion_para.style.font.size = Pt(9)
    conclusion_para.paragraph_format.space_after = Pt(4)
    
    # Small space
    space_para = doc.add_paragraph()
    space_para.paragraph_format.space_after = Pt(4)
    # ========================================
    # NOTES SECTION - WITH BORDERS AND REDUCED HEIGHT
    # ========================================
    notes_heading = doc.add_heading("Notes:", level=2)
    notes_heading.style.font.size = Pt(11)
    notes_heading.style.font.bold = True
    notes_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    notes_heading.paragraph_format.space_before = Pt(4)
    notes_heading.paragraph_format.space_after = Pt(2)

    # Create a 2-row, 1-column table for the Notes section with borders
    notes_table = doc.add_table(rows=2, cols=1)
    notes_table.style = 'Table Grid'
    set_table_borders(notes_table)

    # Set smaller height for both rows (0.4 inches instead of 1.0)
    for row in notes_table.rows:
        row.height = Inches(0.4)

    # Parents remarks row
    parents_cell = notes_table.rows[0].cells[0]
    # Clear any existing content
    parents_cell._element.clear_content()
    # Add the text
    p1 = parents_cell.add_paragraph()
    run1 = p1.add_run("Parents remarks:")
    run1.bold = True
    run1.font.size = Pt(10)

    # Teachers' remarks row
    teachers_cell = notes_table.rows[1].cells[0]
    # Clear any existing content
    teachers_cell._element.clear_content()
    # Add the text
    p2 = teachers_cell.add_paragraph()
    run2 = p2.add_run("Teachers' remarks:")
    run2.bold = True
    run2.font.size = Pt(10)
    
    # Save document with consistent timestamp
    safe_name = "".join([c for c in student_name if c.isalnum() or c==' ']).strip()
    output_path = f"{target_folder}/{safe_name}_Report.docx"
    
    try:
        doc.save(output_path)
        print(f"✓ Saved: {safe_name}_Report.docx")
    except PermissionError:
        timestamp = int(time.time())
        print(f"⚠️ WARNING: Could not save {safe_name}_Report.docx")
        print(f"   Please close this file in Word and try again.")
        alt_path = f"{target_folder}/{safe_name}_Report_{timestamp}.docx"
        doc.save(alt_path)
        print(f"✓ Saved as: {safe_name}_Report_{timestamp}.docx instead")

# =========================
# MAIN
# =========================
def main():
    print("--- Starting Report Generation ---")
    
    # TEST MODE CONFIGURATION
    TEST_MODE = True  # Set to False to process all students
    MAX_STUDENTS = 5  # Number of students to process in test mode
    
    if not os.path.exists(INPUT_FOLDER):
        print(f"❌ Error: {INPUT_FOLDER} folder not found!")
        return
    
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        print(f"✓ Created {OUTPUT_FOLDER} folder")

    files = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(('.xlsx', '.xls', '.csv')) and not f.startswith('~')]
    
    if not files:
        print(f"❌ Error: No Excel/CSV files found in {INPUT_FOLDER}")
        return

    input_path = os.path.join(INPUT_FOLDER, files[0])
    print(f"📂 Loading file: {files[0]}")
    
    try:
        students, class_avg, internal_class = load_and_process_data(input_path)
    except Exception as e:
        print(f"❌ Error loading file: {e}")
        return
    
    # LIMIT STUDENTS IN TEST MODE
    if TEST_MODE:
        original_count = len(students)
        students = students[:MAX_STUDENTS]
        print(f"🧪 TEST MODE: Processing {len(students)} of {original_count} students")
    
    final_class_name = internal_class if internal_class else os.path.splitext(files[0])[0]
    print(f"📚 Class Name: {final_class_name}")
    print(f"👥 Total Students to Process: {len(students)}")
    print("-" * 50)

    for idx, student in enumerate(students, 1):
        name = student.get('name', 'Unknown')
        print(f"\n[{idx}/{len(students)}] Processing: {name}")
        
        try:
            analysis, meta = generate_llm_analysis(student, name, final_class_name)
            
            if analysis:
                create_word_doc(name, analysis, final_class_name, student, class_avg)
            else:
                print(f"❌ Failed to generate analysis for: {name}")
        
        except Exception as e:
            print(f"❌ Error processing {name}: {e}")
            continue
        
        # Rate limiting between students
        if idx < len(students):
            print(f"⏳ Waiting 60 seconds to avoid rate limits...")
            time.sleep(60)
    
    print("\n" + "=" * 50)
    print("✅ Report generation complete!")
    print(f"📁 Output saved to: {OUTPUT_FOLDER}/")

if __name__ == "__main__":
    main()