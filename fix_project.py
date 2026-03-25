import os

# --- CONTENT FOR SRC/DATA_LOADER.PY ---
data_loader_code = r'''import pandas as pd
import numpy as np

def load_and_process_data(file_path):
    """
    Parses the student Excel/CSV file and returns student data and class averages.
    """
    try:
        df = pd.read_csv(file_path)
    except:
        # Fallback if it is actually an excel file renamed as csv
        try:
            df = pd.read_excel(file_path)
        except Exception as e:
            print(f"CRITICAL ERROR: Could not read file. {e}")
            return [], {}

    students = []
    all_scores = {} 

    for index, row in df.iterrows():
        raw_name = row.iloc[0]
        activity = row.iloc[1]
        
        # Identify valid student rows
        if pd.notna(raw_name) and str(raw_name).lower() != 'student name / roll number':
            students.append({
                'name': raw_name,
                'activities': {}
            })
        
        # Process Activity Data
        if students and pd.notna(activity):
            current_student = students[-1]
            scores = []
            feedbacks = []
            
            # Extract Weeks 1-30 (Iterate through columns 2 to ~62)
            # Using safe bounds
            max_col = len(row)
            for i in range(2, min(62, max_col), 2):
                score_val = row.iloc[i]
                feedback_val = row.iloc[i+1] if (i+1) < max_col else None
                
                if pd.notna(score_val):
                    try:
                        s = float(score_val)
                        scores.append(s)
                        if activity not in all_scores:
                            all_scores[activity] = []
                        all_scores[activity].append(s)
                    except:
                        pass
                
                if pd.notna(feedback_val):
                    feedbacks.append(str(feedback_val))
            
            # Calculate Student Average
            student_avg = np.mean(scores) if scores else 0
            
            current_student['activities'][activity] = {
                'scores': scores,
                'feedbacks': feedbacks,
                'average': student_avg
            }

    # Calculate Class Averages
    class_averages = {k: np.mean(v) for k, v in all_scores.items()}
    
    return students, class_averages
'''

# --- CONTENT FOR SRC/GRAPH_GENERATOR.PY ---
graph_code = r'''import matplotlib.pyplot as plt
import numpy as np
import os

def create_radar_chart(student_name, student_data, class_averages, output_dir):
    labels = list(student_data.keys())
    if not labels: return None

    student_scores = [student_data[l]['average'] for l in labels]
    class_scores = [class_averages.get(l, 0) for l in labels]

    num_vars = len(labels)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()

    student_scores += [student_scores[0]]
    class_scores += [class_scores[0]]
    angles += [angles[0]]
    labels += [labels[0]]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    plt.xticks(angles[:-1], labels[:-1])
    ax.set_rlabel_position(0)
    plt.yticks([20, 40, 60, 80], ["20", "40", "60", "80"], color="grey", size=7)
    plt.ylim(0, 100)

    ax.plot(angles, class_scores, color='red', linewidth=2, linestyle='solid', label='Class Avg')
    ax.fill(angles, class_scores, color='red', alpha=0.1)

    ax.plot(angles, student_scores, color='blue', linewidth=2, linestyle='solid', label='Student')
    ax.fill(angles, student_scores, color='blue', alpha=0.25)

    plt.title(f"Profile: {student_name}", size=15, color='black', y=1.1)
    plt.legend(loc='upper right', bbox_to_anchor=(1.1, 1.1))

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    safe_name = "".join([c for c in student_name if c.isalnum() or c==' ']).strip()
    filename = f"{output_dir}/{safe_name}_radar.png"
    plt.savefig(filename)
    plt.close()
    return filename
'''

# --- CONTENT FOR MAIN.PY ---
# Note: You still need to paste your API Key in the file after running this, 
# or paste it below where it says YOUR_KEY_HERE before running.
main_code = r'''import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import time
import json
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.data_loader import load_and_process_data
from src.graph_generator import create_radar_chart

# --- PASTE API KEY BELOW ---
GENAI_API_KEY = "PASTE_YOUR_KEY_HERE"
# ---------------------------

INPUT_FILE = "input/8 CBSC.xlsx - Sheet1.csv"
OUTPUT_FOLDER = "output"

genai.configure(api_key=GENAI_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash') 

def generate_llm_analysis(student_data):
    student_summary = f"Student Name: {student_data['name']}\n"
    for activity, details in student_data['activities'].items():
        avg = round(details['average'])
        feedbacks = "; ".join(details['feedbacks'])
        student_summary += f"- {activity}: Average Score {avg}%. Feedbacks: {feedbacks}\n"

    prompt = f"""
    You are an expert Educational Data Analyst.
    Analyze the following student data and generate content for a report card.
    
    Student Data:
    {student_summary}

    INSTRUCTIONS:
    1. For each activity, determine the "Interest Level" (76-100%: Strong, 51-75%: Positive, 26-50%: Developing, 0-25%: Exploring).
    2. Write a "Key Observation" for each activity (1 sentence, professional).
    3. Suggest 2 "Strong Interest Areas" with potential careers.
    
    OUTPUT FORMAT (Strict JSON):
    {{
      "observations": [
        {{"activity": "Name", "score": 0, "level": "Level", "text": "Observation text"}}
      ],
      "recommendations": [
        {{"area": "Name", "careers": "Career1, Career2", "activity": "Activity suggestion"}}
      ],
      "conclusion": "A short encouraging conclusion sentence."
    }}
    Return ONLY valid JSON.
    """
    
    try:
        response = model.generate_content(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception as e:
        print(f"Error calling Gemini: {e}")
        return None

def create_word_doc(student_name, analysis, graph_path):
    doc = Document()
    heading = doc.add_heading(f"Report Card: {student_name}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if graph_path and os.path.exists(graph_path):
        doc.add_picture(graph_path, width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("Key Observations", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Area', '%', 'Level', 'Observation'

    if analysis:
        for obs in sorted(analysis.get('observations', []), key=lambda x: x['score'], reverse=True):
            row = table.add_row().cells
            row[0].text = obs.get('activity', '')
            row[1].text = f"{obs.get('score', 0)}%"
            row[2].text = obs.get('level', '')
            row[3].text = obs.get('text', '')

        doc.add_heading("Recommendations", level=1)
        rec_table = doc.add_table(rows=1, cols=3)
        rec_table.style = 'Table Grid'
        rec_hdr = rec_table.rows[0].cells
        rec_hdr[0].text, rec_hdr[1].text, rec_hdr[2].text = 'Area', 'Careers', 'Activities'

        for rec in analysis.get('recommendations', []):
            row = rec_table.add_row().cells
            row[0].text = rec.get('area', '')
            row[1].text = rec.get('careers', '')
            row[2].text = rec.get('activity', '')

        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(analysis.get('conclusion', ''))
    
    safe_name = "".join([c for c in student_name if c.isalnum() or c==' ']).strip()
    path = f"{OUTPUT_FOLDER}/{safe_name}_Report.docx"
    doc.save(path)
    print(f"Report saved: {path}")

def main():
    print("--- Starting Report Generation ---")
    if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)

    # Validate file existence
    if not os.path.exists(INPUT_FILE):
        print(f"ERROR: Input file not found at: {INPUT_FILE}")
        return

    students, class_avg = load_and_process_data(INPUT_FILE)
    if not students:
        print("No students found. Check input file format.")
        return
        
    print(f"Loaded {len(students)} students.")

    for student in students:
        name = student['name']
        print(f"Processing: {name}...")
        graph_file = create_radar_chart(name, student['activities'], class_avg, OUTPUT_FOLDER)
        analysis = generate_llm_analysis(student)
        if analysis:
            create_word_doc(name, analysis, graph_file)
        time.sleep(1)

    print("--- Done ---")

if __name__ == "__main__":
    main()
'''

def write_file(path, content):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"✅ Successfully wrote: {path}")

# --- EXECUTION ---
if __name__ == "__main__":
    # Create directories
    os.makedirs("src", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    os.makedirs("input", exist_ok=True)

    # Write files
    write_file("src/data_loader.py", data_loader_code)
    write_file("src/graph_generator.py", graph_code)
    write_file("main.py", main_code)
    
    print("\n SETUP COMPLETE!")
    print("1. Open 'main.py' and PASTE YOUR API KEY.")
    print("2. Run 'python main.py' to generate reports.")